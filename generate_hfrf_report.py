"""
This script is to create the word report file from excle file
"""
import os
import sys
import shutil
import calendar
import datetime
from docx import Document



def get_report_month_input():
    """
    get report month from user to write into file
    """
    try:
        report_month_str = input("Please enter the report month and year (Example: 6/2024): ")
        report_month_split= report_month_str.strip().split("/")
        report_month = int(report_month_split[0])
        report_year = int(report_month_split[1])
    except ValueError as e:
        print ("Input Expect two slash separated and base 10 values, ", e)
        sys.exit()


    def validate_input(report_month: int, report_year: int):
        """internal to the report month validation"""
        months_list = [1,2,3,4,5,6,7,8,9,10,11,12]
        year_list = [2023, 2024, 2025, 2026]
        return (report_month in months_list and report_year in year_list)

    if validate_input(report_month, report_year) is False:
        print("Invalid Input Report Month and Year")
        sys.exit()

    last_day_of_month = calendar._monthlen(report_year, report_month)
    print("Month: ", report_month, ", Year: ", report_year, "last_day_of_month: ",last_day_of_month)
    # return last_day_of_month, report_month, report_year
    return datetime.date(report_year, report_month, last_day_of_month)


def create_new_word_file(base_path):
    """
    from current base path create new working word file 
    by copying and renameing template word file
    """
    template_file = 'template_HFRF_Report.docx'
    report_file_name = 'HFRF_Report.docx'
    template_file_path = os.path.join(base_path, template_file)
    new_file_path = os.path.join(base_path, report_file_name)
    try:
        path = shutil.copyfile(template_file_path,new_file_path)
        return path
    except FileNotFoundError:
        print (f"Report word template file not found at path: {template_file_path}")


def write_into_file(word_file_path, report_month_last_date):
    """start working on word file"""
    document = Document(word_file_path)
    # document.add_heading('A simple text')

    str1 = """
    Monthly Investment Report
    """
    dd = """
    January 2024
    """
    str2 ="""
    Courmacs Legal Ltd
    Alexander House
    Haslingden Road
    Blackburn
    BB1 2EE
    0330 341 0481
    GD@courmacslegal.co.uk
    Courmacslegal.co.uk

    """
    document.add_paragraph(str1)
    document.add_paragraph(dd)
    document.add_paragraph(str2)

    document.save(word_file_path)


def main():
    """This is main function"""
    report_month_last_date = get_report_month_input()
    base_path = os.getcwd()
    new_file_path = create_new_word_file(base_path)
    # print (new_file_path)
    write_into_file(new_file_path, report_month_last_date)


if __name__ == "__main__":
    main()
